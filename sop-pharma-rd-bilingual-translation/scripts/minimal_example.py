#!/usr/bin/env python3
from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


def style_english_run(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    para.add_run(text)
    return para


def write_bilingual_title(paragraph, english: str, chinese: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(f"{english} {chinese}")
    run.bold = True
    style_english_run(run)


def write_bilingual_paragraph(paragraph, english: str, chinese: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(english)
    style_english_run(run)
    insert_paragraph_after(paragraph, chinese)


def write_bilingual_cell(cell, english_paragraphs: list[str], chinese_paragraphs: list[str]) -> None:
    cell.text = ""

    first = True
    for english in english_paragraphs:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = para.add_run(english)
        style_english_run(run)

    for chinese in chinese_paragraphs:
        cell.add_paragraph(chinese)


def translate_example(source: Path) -> tuple[Path, Path]:
    output = source.with_name(f"{source.stem}简约示例翻译版{source.suffix}")
    report = source.with_name(f"{source.stem}简约示例_QA.md")

    shutil.copy2(source, output)
    doc = Document(output)

    # 1) 标题示例：单行双语标题
    title_map = {
        "目的": "Purpose",
        "范围": "Scope",
    }

    # 2) 正文段落示例：英文段在前，中文段在后
    paragraph_map = {
        "本指导书适用于先声临床统计与数据管理室所有参与TEAE分析处理流程的员工，包括外包项目合作方相关人员。":
            "This WI applies to all personnel involved in the TEAE analysis workflow, including relevant personnel from outsourced project partners.",
        "该指南的不良事件收集方式均提供一种思路，非强制，各个项目最终收集方式以项目组成员根据本试验特殊具体情况而决定。":
            "The adverse event collection approach in this guideline is provided as a reference and is not mandatory; the final collection approach for each project shall be decided by the project team according to the specific trial circumstances.",
    }

    # 3) 表格示例：多行文本单元格，英文块在前，中文块在后
    cell_map = {
        "不良事件转归": "Adverse Event Outcome",
        "○痊愈后无后遗症": "Recovered without sequelae",
        "○痊愈后有后遗症": "Recovered with sequelae",
        "○好转": "Improved",
        "○未愈": "Not recovered",
        "○死亡": "Death",
        "○未知": "Unknown",
        "严重程度（根据方案选择）": "Severity (Per Protocol)",
        "○ 1级": "Grade 1",
        "○ 2级": "Grade 2",
        "○ 3级": "Grade 3",
        "○ 4级": "Grade 4",
        "○ 5级": "Grade 5",
    }

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if text in title_map:
            write_bilingual_title(paragraph, title_map[text], text)
        elif text in paragraph_map:
            write_bilingual_paragraph(paragraph, paragraph_map[text], text)

    # 跳过第一张审批表，只演示正文表格处理
    for table in doc.tables[1:]:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                chinese_paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                if not chinese_paragraphs:
                    continue

                if row_index == 0 and len(chinese_paragraphs) == 1:
                    chinese = chinese_paragraphs[0]
                    english = cell_map.get(chinese, chinese)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(f"{english} {chinese}")
                    style_english_run(run)
                    continue

                english_paragraphs = [cell_map.get(text, text) for text in chinese_paragraphs]
                if english_paragraphs != chinese_paragraphs:
                    write_bilingual_cell(cell, english_paragraphs, chinese_paragraphs)

    doc.save(output)
    report.write_text(
        "\n".join(
            [
                "# Minimal QA",
                "",
                f"- Source: {source}",
                f"- Output: {output}",
                "- Included examples:",
                "  - bilingual title",
                "  - bilingual paragraph",
                "  - multi-line table cell with English block first",
            ]
        ),
        encoding="utf-8",
    )
    return output, report


def main() -> None:
    parser = argparse.ArgumentParser(description="Minimal bilingual SOP translation example.")
    parser.add_argument("source", help="Path to source .docx file")
    args = parser.parse_args()

    source = Path(args.source).expanduser().resolve()
    if not source.exists() or source.suffix.lower() != ".docx":
        raise SystemExit(f"Invalid source file: {source}")

    output, report = translate_example(source)
    print(f"output={output}")
    print(f"qa={report}")


if __name__ == "__main__":
    main()
